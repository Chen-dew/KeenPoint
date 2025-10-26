"""
文档解析服务
处理 PDF 和 Word 文档的解析与文本提取
使用 MinerU API 进行 PDF 解析
"""

from docx import Document
import tempfile
import os
import zipfile
import aiohttp
import aiofiles
from typing import Dict
from app.core.logger import logger
from app.core.utils import generate_unique_id
from app.core.config import settings

async def parse_pdf_with_mineru(pdf_path: str, output_folder: str) -> str:
    """
    使用 MinerU API 解析 PDF 文件并提取文本和图像
    
    参数:
    - pdf_path: PDF 文件路径
    - output_folder: 保存提取内容的根目录
    
    返回:
    - str: 提取内容的文件夹路径
    """
    assert settings.MINERU_API is not None, "MINERU_API is not set"
    logger.info(f"🔧 使用 MinerU API 解析 PDF: {pdf_path}")
    
    os.makedirs(output_folder, exist_ok=True)

    # 读取 PDF 文件内容
    async with aiofiles.open(pdf_path, "rb") as f:
        pdf_content = await f.read()

    # 准备表单数据
    data = aiohttp.FormData()
    data.add_field(
        "files",
        pdf_content,
        filename=os.path.basename(pdf_path),
        content_type="application/pdf",
    )
    data.add_field("return_images", "True")
    data.add_field("response_format_zip", "True")

    # 准备请求头
    headers = {
        "Authorization": f"Bearer {settings.MINERU_TOKEN}"
    }

    # 发送请求到 MinerU API
    async with aiohttp.ClientSession() as session:
        try:
            async with session.post(settings.MINERU_API, data=data, headers=headers) as response:
                response.raise_for_status()
                content = await response.read()

                # 保存 ZIP 文件
                with tempfile.NamedTemporaryFile(delete=False, suffix=".zip") as tmp:
                    tmp.write(content)
                    zip_path = tmp.name

                logger.info(f"📦 收到 ZIP 响应，正在解压...")

                # 解压 ZIP 文件
                with zipfile.ZipFile(zip_path, "r") as zip_ref:
                    # 获取顶层文件夹名称
                    top_level = {
                        name.split("/", 1)[0] for name in zip_ref.namelist() if name.strip()
                    }
                    if len(top_level) != 1:
                        raise RuntimeError("Expected exactly one top-level folder in zip")
                    prefix = list(top_level)[0] + "/"

                    # 提取所有文件
                    for member in zip_ref.infolist():
                        filename = member.filename
                        dest_path = os.path.join(
                            output_folder, filename.removeprefix(prefix)
                        )

                        if not member.is_dir():
                            os.makedirs(os.path.dirname(dest_path), exist_ok=True)
                            with zip_ref.open(member) as src, open(dest_path, "wb") as dst:
                                dst.write(src.read())

                # 清理临时 ZIP 文件
                try:
                    os.unlink(zip_path)
                except:
                    pass

                logger.info(f"✅ PDF 解析完成，输出目录: {output_folder}")
                return output_folder

        except aiohttp.ClientError as e:
            logger.error(f"❌ MinerU API 请求失败: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"❌ PDF 解析失败: {str(e)}")
            raise


async def parse_document(file) -> Dict:
    """
    解析 PDF 或 Word 文件
    
    参数:
    - file: UploadFile 对象
    
    返回:
    - 解析结果字典
    """
    suffix = file.filename.split(".")[-1].lower()
    
    if suffix == "pdf":
        return await _parse_pdf_advanced(file)
    elif suffix in ["doc", "docx"]:
        return await _parse_word(file)
    else:
        logger.error(f"不支持的文件类型: {suffix}")
        return {"error": "Unsupported file type"}

async def _parse_pdf_advanced(file) -> Dict:
    """
    使用 MinerU API 解析 PDF 文件
    
    参数:
    - file: UploadFile 对象
    
    返回:
    - PDF 解析结果
    """
    try:
        # 保存上传的文件到临时位置
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name
        
        # 创建输出文件夹
        document_id = generate_unique_id()
        output_folder = os.path.join(settings.OUTPUT_DIR, "parsed", document_id)
        
        # 使用 MinerU API 解析
        extracted_folder = await parse_pdf_with_mineru(tmp_path, output_folder)
        
        # 读取解析后的 Markdown 文件
        md_files = [f for f in os.listdir(extracted_folder) if f.endswith(".md")]
        full_text = ""
        
        if md_files:
            md_path = os.path.join(extracted_folder, md_files[0])
            async with aiofiles.open(md_path, "r", encoding="utf-8") as f:
                full_text = await f.read()
        
        # 统计图像数量
        images_folder = os.path.join(extracted_folder, "images")
        image_count = 0
        if os.path.exists(images_folder):
            image_count = len([f for f in os.listdir(images_folder) if f.lower().endswith(('.png', '.jpg', '.jpeg'))])
        
        # 统计页数（通过 Markdown 文件或图像数量估算）
        page_count = 0
        if md_files:
            # 简单估算：根据文本长度估算页数
            page_count = max(1, len(full_text) // 3000)
        
        # 清理临时文件
        try:
            os.unlink(tmp_path)
        except:
            pass
        
        result = {
            "type": "pdf",
            "filename": file.filename,
            "page_count": page_count,
            "text_length": len(full_text),
            "text_preview": full_text[:1000] + "..." if len(full_text) > 1000 else full_text,
            "full_text": full_text,
            "markdown_path": os.path.join(extracted_folder, md_files[0]) if md_files else None,
            "images_folder": images_folder if os.path.exists(images_folder) else None,
            "image_count": image_count,
            "output_folder": extracted_folder,
            "metadata": {
                "title": "",
                "author": "",
                "subject": "",
                "creator": "MinerU API"
            },
            "document_id": document_id,
            "parsing_method": "mineru_api"
        }
        
        logger.info(f"✅ PDF 解析成功 (MinerU): {file.filename}, {page_count} 页估算, {image_count} 张图像")
        return result
    
    except Exception as e:
        logger.error(f"❌ MinerU PDF 解析失败: {str(e)}")
        return {"error": f"PDF parsing failed: {str(e)}"}

async def _parse_word(file) -> Dict:
    """
    解析 Word 文件
    
    参数:
    - file: UploadFile 对象
    
    返回:
    - Word 解析结果
    """
    try:
        # 创建临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            content = await file.read()
            tmp.write(content)
            tmp_path = tmp.name
        
        # 使用 python-docx 解析
        doc = Document(tmp_path)
        
        # 提取段落文本
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
        
        # 提取表格内容
        tables_count = len(doc.tables)
        
        # 提取核心属性
        core_properties = doc.core_properties
        
        result = {
            "type": "word",
            "filename": file.filename,
            "paragraph_count": len(paragraphs),
            "tables_count": tables_count,
            "text_length": len(text),
            "text_preview": text[:1000] + "..." if len(text) > 1000 else text,
            "full_text": text,
            "metadata": {
                "title": core_properties.title or "",
                "author": core_properties.author or "",
                "subject": core_properties.subject or "",
                "keywords": core_properties.keywords or ""
            },
            "document_id": generate_unique_id()
        }
        
        # 清理临时文件
        try:
            os.unlink(tmp_path)
        except:
            pass
        
        logger.info(f"✅ Word 解析成功: {file.filename}, {len(paragraphs)} 段落")
        return result
    
    except Exception as e:
        logger.error(f"❌ Word 解析失败: {str(e)}")
        return {"error": f"Word parsing failed: {str(e)}"}

async def extract_text_from_file(file_path: str) -> str:
    """
    从文件路径直接提取文本
    
    参数:
    - file_path: 文件路径
    
    返回:
    - 提取的文本
    """
    ext = file_path.split('.')[-1].lower()
    
    try:
        if ext == 'pdf':
            # 使用 MinerU API 解析 PDF
            document_id = generate_unique_id()
            output_folder = os.path.join(settings.OUTPUT_DIR, "temp_extract", document_id)
            
            extracted_folder = await parse_pdf_with_mineru(file_path, output_folder)
            
            # 读取 Markdown 文件
            md_files = [f for f in os.listdir(extracted_folder) if f.endswith(".md")]
            if md_files:
                md_path = os.path.join(extracted_folder, md_files[0])
                async with aiofiles.open(md_path, "r", encoding="utf-8") as f:
                    return await f.read()
            return ""
            
        elif ext in ['doc', 'docx']:
            doc = Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs)
        else:
            return ""
    except Exception as e:
        logger.error(f"❌ 文本提取失败: {str(e)}")
        return ""
